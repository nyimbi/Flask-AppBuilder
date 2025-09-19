import hashlib
import logging
import mimetypes
import os
from io import BytesIO
from typing import Any, Dict, List, Optional

import docx
import magic
import markdown
import nltk
import pdfkit
import pypandoc
import PyPDF2
from cryptography.fernet import Fernet
from flask_appbuilder.models.mixins import ImageColumn
from nltk.corpus import stopwords
from nltk.tokenize import sent_tokenize, word_tokenize
from sqlalchemy import Boolean, Column, FileColumn, Integer, LargeBinary, String, Text
from sqlalchemy.dialects.postgresql import ARRAY, JSONB
from sqlalchemy.ext.declarative import declared_attr
from sqlalchemy.orm import validates
from sqlalchemy.sql import func, text
from sqlalchemy_utils import TSVectorType
from transformers import pipeline

from base_mixin import BaseModelMixin

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocMixin(BaseModelMixin):
    """
    A mixin class for document management in SQLAlchemy models.

    This mixin provides extensive functionality for handling various types of documents,
    including metadata extraction, format conversion, encryption, full-text search,
    and more.

    Attributes:
        mime_type (str): The MIME type of the document.
        doc (ImageColumn): An image representation of the document.
        doc_text (str): The text content of the document.
        doc_binary (FileColumn): The binary content of the document.
        doc_title (str): The title of the document.
        subject (str): The subject of the document.
        author (str): The author of the document.
        keywords (str): Keywords associated with the document.
        comments (str): Comments on the document.
        chapter_number (int): The chapter number in the document structure.
        chapter_title (str): The title of the chapter.
        section_number (int): The section number in the document structure.
        section_title (str): The title of the section.
        sub_section_number (int): The sub-section number in the document structure.
        sub_section_title (str): The title of the sub-section.
        doc_context (str): Context for LLM-based text generation.
        doc_prompt (str): Prompt for LLM-based text generation.
        doc_type (str): The type of the document (e.g., 'pdf', 'docx').
        char_count (int): The character count of the document.
        word_count (int): The word count of the document.
        lines (int): The number of lines in the document.
        paragraphs (int): The number of paragraphs in the document.
        gpt_token_count (int): The token count for GPT models.
        grammar_checked (bool): Whether the document has been grammar checked.
        doc_summary (str): A summary of the document.
        doc_spell_checked (bool): Whether the document has been spell-checked.
        doc_gpt_ver (str): The version of GPT used for processing.
        doc_format (str): The format of the document.
        doc_downloadable (bool): Whether the document is downloadable.
        doc_template (str): The template used for the document.
        doc_rendered (bool): Whether the document has been rendered.
        doc_render (FileColumn): The rendered version of the document.
        file_size_bytes (int): The size of the document in bytes.
        producer_prog (str): The program that produced the document.
        immutable (bool): Whether the document is immutable.
        page_size (str): The page size of the document.
        page_count (int): The number of pages in the document.
        hashx (str): A hash of the document content.
        is_audio (bool): Whether the document is an audio file.
        audio_duration_secs (int): The duration of the audio in seconds.
        audio_frame_rate (int): The frame rate of the audio.
        audio_channels (int): The number of audio channels.
        is_encrypted (bool): Whether the document is encrypted.
        encryption_key (str): The encryption key for the document.
    """

    mime_type: str = Column(String(60), default="application/pdf", nullable=False)
    doc: ImageColumn = Column(
        ImageColumn(thumbnail_size=(30, 30, True), size=(300, 300, True))
    )
    doc_text: str = Column(Text, nullable=True)
    doc_binary: FileColumn = Column(FileColumn, nullable=False)
    doc_title: str = Column(String(200), nullable=False, default="Untitled Document")
    subject: str = Column(String(100), nullable=True)
    author: str = Column(String(100), nullable=True)
    keywords: str = Column(String(200), nullable=True)
    comments: str = Column(Text, nullable=True)

    # fields for chapter structure
    chapter_number: int = Column(Integer, nullable=True)
    chapter_title: str = Column(String(200), nullable=True)
    section_number: int = Column(Integer, nullable=True)
    section_title: str = Column(String(200), nullable=True)
    sub_section_number: int = Column(Integer, nullable=True)
    sub_section_title: str = Column(String(200), nullable=True)

    # fields for LLM-based text generation
    doc_context: str = Column(Text, nullable=True)
    doc_prompt: str = Column(Text, nullable=True)

    # Metadata
    doc_type: str = Column(String(5), default="pdf", nullable=False)
    char_count: int = Column(Integer, default=0)
    word_count: int = Column(Integer, default=0)
    lines: int = Column(Integer, default=0)
    paragraphs: int = Column(Integer, default=0)
    gpt_token_count: int = Column(Integer, default=0)
    grammar_checked: bool = Column(Boolean, default=False)
    doc_summary: str = Column(Text, nullable=True)
    doc_spell_checked: bool = Column(Boolean, default=False)
    doc_gpt_ver: str = Column(String(40), nullable=True)
    doc_format: str = Column(String(40), nullable=False, default="pdf")
    doc_downloadable: bool = Column(Boolean, default=True)
    doc_template: str = Column(Text, nullable=True)
    doc_rendered: bool = Column(Boolean, default=False)
    doc_render: FileColumn = Column(FileColumn, nullable=True)

    file_size_bytes: int = Column(Integer, default=0)
    producer_prog: str = Column(String(40), nullable=True)
    immutable: bool = Column(Boolean, default=False)

    page_size: str = Column(String(40), nullable=True)
    page_count: int = Column(Integer, default=1)
    hashx: str = Column(String(64), nullable=False)  # SHA-256 hash

    # Audio metadata
    is_audio: bool = Column(Boolean, default=False)
    audio_duration_secs: int = Column(Integer, nullable=True)
    audio_frame_rate: int = Column(Integer, nullable=True)
    audio_channels: int = Column(Integer, nullable=True)

    # Encryption
    is_encrypted: bool = Column(Boolean, default=False)
    encryption_key: str = Column(String(100), nullable=True)

    # PostgreSQL-specific columns
    metadata: dict = Column(JSONB, nullable=True)
    versions: list = Column(ARRAY(JSONB), nullable=True)
    tags: list = Column(ARRAY(String), nullable=True)

    @classmethod
    def check_dependencies(cls) -> Dict[str, bool]:
        """
        Check if all required dependencies are installed.

        Returns:
            Dict[str, bool]: A dictionary with dependency names as keys and boolean values
                             indicating whether they are installed.
        """
        dependencies = {
            "PyPDF2": True,
            "pypandoc": True,
            "python-magic": True,
            "python-docx": True,
            "nltk": True,
            "transformers": True,
            "markdown": True,
            "pdfkit": True,
        }

        try:
            import PyPDF2
        except ImportError:
            dependencies["PyPDF2"] = False

        try:
            import pypandoc
        except ImportError:
            dependencies["pypandoc"] = False

        try:
            import magic
        except ImportError:
            dependencies["python-magic"] = False

        try:
            import docx
        except ImportError:
            dependencies["python-docx"] = False

        try:
            import nltk
        except ImportError:
            dependencies["nltk"] = False

        try:
            from transformers import pipeline
        except ImportError:
            dependencies["transformers"] = False

        try:
            import markdown
        except ImportError:
            dependencies["markdown"] = False

        try:
            import pdfkit
        except ImportError:
            dependencies["pdfkit"] = False

        return dependencies

    def _extract_pdf_metadata(self) -> Dict[str, Any]:
        """Extract metadata from PDF files."""
        metadata = {}
        try:
            with BytesIO(self.doc_binary) as file:
                pdf = PyPDF2.PdfReader(file)
                info = pdf.metadata
                if info:
                    metadata.update(
                        {
                            "author": info.get("/Author", ""),
                            "subject": info.get("/Subject", ""),
                            "title": info.get("/Title", ""),
                            "creator": info.get("/Creator", ""),
                            "producer": info.get("/Producer", ""),
                            "creation_date": info.get("/CreationDate", ""),
                            "modification_date": info.get("/ModDate", ""),
                            "keywords": info.get("/Keywords", ""),
                        }
                    )
                metadata["page_count"] = len(pdf.pages)
                metadata["first_page_text"] = pdf.pages[0].extract_text()[:1000]
        except Exception as e:
            logger.error(f"Error extracting PDF metadata: {str(e)}")
            raise
        return metadata

    def _extract_word_metadata(self) -> Dict[str, Any]:
        """Extract metadata from Word documents."""
        metadata = {}
        try:
            with BytesIO(self.doc_binary) as file:
                doc = docx.Document(file)
                core_props = doc.core_properties
                metadata.update(
                    {
                        "author": core_props.author or "",
                        "title": core_props.title or "",
                        "subject": core_props.subject or "",
                        "keywords": core_props.keywords or "",
                        "created": (
                            str(core_props.created) if core_props.created else ""
                        ),
                        "modified": (
                            str(core_props.modified) if core_props.modified else ""
                        ),
                        "last_modified_by": core_props.last_modified_by or "",
                        "revision": core_props.revision or 1,
                        "category": core_props.category or "",
                        "paragraphs": len(doc.paragraphs),
                        "sections": len(doc.sections),
                    }
                )
        except Exception as e:
            logger.error(f"Error extracting Word metadata: {str(e)}")
            metadata["extraction_error"] = str(e)
        return metadata

    def _extract_text_metadata(self) -> Dict[str, Any]:
        """Extract metadata from plain text files."""
        content = self.doc_binary.decode("utf-8", errors="ignore")
        lines = content.split("\n")
        return {
            "file_size": len(content),
            "line_count": len(lines),
            "first_line": lines[0][:100] if lines else "",
            "encoding": "utf-8",
            "has_bom": content.startswith("\ufeff"),
            "empty_lines": len([l for l in lines if not l.strip()]),
            "avg_line_length": sum(len(l) for l in lines) / len(lines) if lines else 0,
        }

    def _extract_markdown_metadata(self) -> Dict[str, Any]:
        """Extract metadata from markdown files."""
        content = self.doc_binary.decode("utf-8", errors="ignore")
        lines = content.split("\n")
        headers = [line for line in lines if line.strip().startswith("#")]
        return {
            "file_size": len(content),
            "headers_count": len(headers),
            "first_heading": headers[0][:100] if headers else "",
            "links_count": content.count("]("),
            "code_blocks": content.count("```"),
            "images": content.count("!["),
            "lists": len(
                [l for l in lines if l.strip().startswith(("- ", "* ", "1. "))]
            ),
        }

    def _update_attributes_from_metadata(self, metadata: Dict[str, Any]) -> None:
        """Update object attributes with extracted metadata."""
        for key, value in metadata.items():
            if hasattr(self, key) and value is not None:
                setattr(self, key, value)

        # Store complete metadata in JSONB field
        if self.metadata is None:
            self.metadata = {}
        self.metadata.update(metadata)

        # Update hash
        self.update_hash()

    def update_hash(self) -> None:
        """Update document hash using SHA-256."""
        if self.doc_binary:
            self.hashx = hashlib.sha256(self.doc_binary).hexdigest()

    @declared_attr
    def search_vector(cls):
        return Column(
            TSVectorType(
                "doc_title",
                "doc_text",
                "comments",
                weights={"doc_title": "A", "doc_text": "B", "comments": "C"},
            ),
            nullable=False,
            index=True,
        )

    @validates("doc_text")
    def update_doc_info(self, key: str, value: str) -> str:
        if value:
            self.char_count = len(value)
            self.word_count = len(value.split())
            self.lines = value.count("\n") + 1
            self.paragraphs = len([p for p in value.split("\n\n") if p.strip()])
        return value

    def extract_metadata(self) -> Dict[str, Any]:
        """
        Extract metadata from the document and populate relevant fields.

        :return: A dictionary containing the extracted metadata.
        """
        metadata = {}

        mime_type_handlers = {
            "application/pdf": self._extract_pdf_metadata,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": self._extract_word_metadata,
            "application/msword": self._extract_word_metadata,
            "text/plain": self._extract_text_metadata,
            "text/markdown": self._extract_markdown_metadata,
        }

        handler = mime_type_handlers.get(self.mime_type)
        if handler:
            try:
                metadata = handler()
            except Exception as e:
                metadata["extraction_error"] = str(e)
                logger.error(f"Error extracting metadata: {str(e)}")
        else:
            metadata["extraction_note"] = (
                f"Metadata extraction not supported for MIME type: {self.mime_type}"
            )

        self._update_attributes_from_metadata(metadata)
        return metadata

    def generate_summary(self, max_length: int = 200) -> None:
        """
        Generate a summary of the document content using NLP techniques.

        Args:
            max_length: Maximum length of the summary in words
        """
        try:
            nltk.download("punkt", quiet=True)
            nltk.download("stopwords", quiet=True)

            if not self.doc_text:
                raise ValueError("No document text available for summarization")

            # Tokenize the text into sentences
            sentences = sent_tokenize(self.doc_text)
            if not sentences:
                raise ValueError("No sentences found in document text")

            # Calculate word frequencies
            stop_words = set(stopwords.words("english"))
            words = word_tokenize(self.doc_text.lower())
            word_frequencies = {}

            for word in words:
                if word not in stop_words and word.isalnum():
                    word_frequencies[word] = word_frequencies.get(word, 0) + 1

            # Normalize frequencies
            max_frequency = max(word_frequencies.values()) if word_frequencies else 1
            for word in word_frequencies.keys():
                word_frequencies[word] = word_frequencies[word] / max_frequency

            # Calculate sentence scores
            sentence_scores = {}
            for sentence in sentences:
                if len(sentence.split()) < 30:  # Skip very long sentences
                    score = 0
                    for word in word_tokenize(sentence.lower()):
                        if word in word_frequencies:
                            score += word_frequencies[word]
                    sentence_scores[sentence] = score

            # Get top sentences
            summary_sentences = sorted(
                sentence_scores.items(), key=lambda x: x[1], reverse=True
            )[:3]

            summary = " ".join(sentence for sentence, score in summary_sentences)

            # Use transformers for advanced summarization if available
            try:
                summarizer = pipeline("summarization", model="facebook/bart-large-cnn")
                chunks = [summary[i : i + 1024] for i in range(0, len(summary), 1024)]
                summarized_chunks = []

                for chunk in chunks:
                    result = summarizer(
                        chunk, max_length=max_length, min_length=30, do_sample=False
                    )
                    summarized_chunks.append(result[0]["summary_text"])

                self.doc_summary = " ".join(summarized_chunks)
            except Exception as e:
                logger.warning(
                    f"Advanced summarization failed, using extractive summary: {str(e)}"
                )
                self.doc_summary = summary

        except Exception as e:
            logger.error(f"Error generating summary: {str(e)}")
            raise

    @classmethod
    def search(
        cls,
        session,
        query: str,
        *,
        limit: Optional[int] = None,
        offset: Optional[int] = None,
    ) -> List["DocMixin"]:
        """
        Perform a full-text search on the doc_title, doc_text and comments fields.

        Args:
            session: A SQLAlchemy session object
            query: The search query
            limit: Maximum number of results to return
            offset: Number of results to skip

        Returns:
            List of matching documents with relevance scores and highlights
        """
        try:
            # Sanitize query
            query = " & ".join(word.strip() for word in query.split() if word.strip())
            search_query = func.plainto_tsquery("english", query)
            rank_function = func.ts_rank(cls.search_vector, search_query)

            # Build query
            results = (
                session.query(cls, rank_function.label("relevance"))
                .filter(cls.search_vector.match(search_query))
                .order_by(rank_function.desc())
            )

            # Apply pagination
            if offset is not None:
                results = results.offset(offset)
            if limit is not None:
                results = results.limit(limit)

            # Process results
            highlighted_results = []
            for doc, relevance in results.all():
                doc.relevance_score = float(relevance)
                doc.highlighted_title = cls.highlight_matched_terms(
                    doc.doc_title or "", query
                )
                doc.highlighted_text = cls.highlight_matched_terms(
                    doc.doc_text or "", query
                )
                highlighted_results.append(doc)

            return highlighted_results

        except Exception as e:
            logger.error(f"Search error: {str(e)}")
            raise

    @staticmethod
    def highlight_matched_terms(text: str, query: str) -> str:
        """
        Highlight matched terms in text using HTML spans.

        Args:
            text: Source text
            query: Search query

        Returns:
            Text with highlighted search terms
        """
        if not text or not query:
            return text

        words = set(word.lower() for word in query.split() if word.strip())
        result = text

        for word in words:
            pattern = re.compile(f"({word})", re.IGNORECASE)
            result = pattern.sub(r'<span class="highlight">\1</span>', result)

        return result

    def encrypt_document(self, key: Optional[bytes] = None) -> None:
        """
        Encrypt document using Fernet symmetric encryption.

        Args:
            key: Optional encryption key (will generate if not provided)
        """
        if self.is_encrypted:
            raise ValueError("Document is already encrypted")

        try:
            if not key:
                key = Fernet.generate_key()

            fernet = Fernet(key)
            self.doc_binary = fernet.encrypt(self.doc_binary)
            self.encryption_key = key.decode("utf-8")
            self.is_encrypted = True

            # Store encryption metadata
            if self.metadata is None:
                self.metadata = {}
            self.metadata["encryption"] = {
                "method": "Fernet",
                "timestamp": datetime.now().isoformat(),
            }

        except Exception as e:
            logger.error(f"Encryption error: {str(e)}")
            raise

    def decrypt_document(self) -> None:
        """Decrypt document using stored encryption key."""
        if not self.is_encrypted:
            raise ValueError("Document is not encrypted")

        if not self.encryption_key:
            raise ValueError("Encryption key not found")

        try:
            fernet = Fernet(self.encryption_key.encode("utf-8"))
            self.doc_binary = fernet.decrypt(self.doc_binary)
            self.is_encrypted = False
            self.encryption_key = None

            if self.metadata and "encryption" in self.metadata:
                del self.metadata["encryption"]

        except Exception as e:
            logger.error(f"Decryption error: {str(e)}")
            raise

    def convert_format(self, target_format: str) -> None:
        """
        Convert document to different format using Pandoc.

        Args:
            target_format: Target format ('pdf', 'docx', 'md', 'html', etc.)
        """
        if self.doc_type == target_format:
            return

        format_mapping = {
            "pdf": "pdf",
            "docx": "docx",
            "md": "markdown",
            "html": "html",
            "txt": "plain",
            "rst": "rst",
            "epub": "epub",
            "odt": "odt",
            "latex": "latex",
        }

        source_format = format_mapping.get(self.doc_type)
        target_pandoc_format = format_mapping.get(target_format)

        if not source_format or not target_pandoc_format:
            raise ValueError(
                f"Conversion from {self.doc_type} to {target_format} not supported"
            )

        try:
            # Create temporary file for input
            with NamedTemporaryFile(
                suffix=f".{self.doc_type}", delete=False
            ) as temp_in:
                temp_in.write(self.doc_binary)
                temp_in.flush()

            # Convert document
            output = pypandoc.convert_file(
                temp_in.name,
                target_pandoc_format,
                format=source_format,
                outputfile=None,
            )

            # Update document properties
            self.doc_binary = (
                output.encode("utf-8") if isinstance(output, str) else output
            )
            self.doc_type = target_format
            self.mime_type = (
                mimetypes.guess_type(f"dummy.{target_format}")[0]
                or "application/octet-stream"
            )

            # Update metadata
            self.extract_metadata()

            # Add conversion to version history
            if self.versions is None:
                self.versions = []
            self.versions.append(
                {
                    "timestamp": datetime.now().isoformat(),
                    "action": "format_conversion",
                    "from_format": self.doc_type,
                    "to_format": target_format,
                }
            )

        except Exception as e:
            logger.error(f"Format conversion error: {str(e)}")
            raise
        finally:
            try:
                os.unlink(temp_in.name)
            except:
                pass

    def to_markdown(self) -> str:
        """Convert document to Markdown format."""
        try:
            self.convert_format("md")
            return self.doc_binary.decode("utf-8")
        except Exception as e:
            logger.error(f"Markdown conversion error: {str(e)}")
            raise

    def from_markdown(self, markdown_text: str, target_format: str) -> None:
        """
        Convert Markdown text to specified format.

        Args:
            markdown_text: Source markdown text
            target_format: Target format to convert to
        """
        try:
            with NamedTemporaryFile(mode="w", suffix=".md", delete=False) as temp_file:
                temp_file.write(markdown_text)
                temp_file.flush()

                output = pypandoc.convert_file(
                    temp_file.name, target_format, format="markdown", outputfile=None
                )

            self.doc_binary = (
                output.encode("utf-8") if isinstance(output, str) else output
            )
            self.doc_type = target_format
            self.mime_type = (
                mimetypes.guess_type(f"dummy.{target_format}")[0]
                or "application/octet-stream"
            )

            self.extract_metadata()

        except Exception as e:
            logger.error(f"Markdown conversion error: {str(e)}")
            raise
        finally:
            try:
                os.unlink(temp_file.name)
            except:
                pass

    def generate_text_with_llm(self, llm_function: Callable[[str, str], str]) -> None:
        """
        Generate text using an LLM based on context and prompt.

        Args:
            llm_function: Function that takes context and prompt and returns generated text
        """
        try:
            if not self.doc_context or not self.doc_prompt:
                raise ValueError("Both context and prompt are required")

            generated_text = llm_function(self.doc_context, self.doc_prompt)

            if not generated_text:
                raise ValueError("LLM returned empty text")

            self.doc_text = generated_text
            self.update_doc_info("doc_text", generated_text)

            # Store generation metadata
            if self.metadata is None:
                self.metadata = {}
            self.metadata["llm_generation"] = {
                "timestamp": datetime.now().isoformat(),
                "context_length": len(self.doc_context),
                "prompt_length": len(self.doc_prompt),
                "output_length": len(generated_text),
            }

        except Exception as e:
            logger.error(f"Text generation error: {str(e)}")
            raise

    def detect_mime_type(self, filename: Optional[str] = None) -> str:
        """
        Detect and set document MIME type.

        Args:
            filename: Optional filename for extension-based detection

        Returns:
            Detected MIME type
        """
        try:
            if not self.doc_binary and not filename:
                raise ValueError("Either doc_binary or filename required")

            if self.doc_binary:
                mime = magic.Magic(mime=True)
                detected_mime_type = mime.from_buffer(self.doc_binary)
            else:
                detected_mime_type, _ = mimetypes.guess_type(filename)

            if not detected_mime_type:
                detected_mime_type = "application/octet-stream"

            self.mime_type = detected_mime_type
            return detected_mime_type

        except Exception as e:
            logger.error(f"MIME type detection error: {str(e)}")
            raise

    def set_doc_type_from_mime_type(self) -> None:
        """Set doc_type based on detected MIME type."""
        mime_to_doc_type = {
            "application/pdf": "pdf",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
            "application/msword": "doc",
            "text/plain": "txt",
            "text/markdown": "md",
            "text/html": "html",
            "application/rtf": "rtf",
            "application/epub+zip": "epub",
            "application/x-latex": "latex",
        }
        self.doc_type = mime_to_doc_type.get(self.mime_type, "unknown")

    def update_document(
        self, file_content: bytes, filename: Optional[str] = None
    ) -> None:
        """
        Update document content and metadata.

        Args:
            file_content: New binary content
            filename: Optional filename for MIME detection
        """
        if self.immutable:
            raise ValueError("Cannot update immutable document")

        try:
            # Store previous version
            if self.versions is None:
                self.versions = []
            self.versions.append(
                {
                    "timestamp": datetime.now().isoformat(),
                    "action": "update",
                    "previous_hash": self.hashx,
                    "previous_size": self.file_size_bytes,
                }
            )

            self.doc_binary = file_content
            self.detect_mime_type(filename)
            self.set_doc_type_from_mime_type()
            self.file_size_bytes = len(file_content)
            self.update_hash()
            self.extract_metadata()

        except Exception as e:
            logger.error(f"Document update error: {str(e)}")
            raise
